# app.py

import os
from flask import Flask, request, render_template, send_file, jsonify
import openai
from PyPDF2 import PdfReader
from docx import Document
import yaml
import uuid

# Load configuration from YAML file
def load_config(config_path='config.yaml'):
    with open(config_path, 'r') as file:
        config = yaml.safe_load(file)
    return config

config = load_config()

app = Flask(__name__)
app.config['SECRET_KEY'] = os.urandom(24)  # For session security, if needed

# Set OpenAI API key and model from config
openai.api_key = config['openai']['api_key']
OPENAI_MODEL = config['openai'].get('model', 'gpt-4')

def extract_text_from_pdf(pdf_file):
    """Extract text from an uploaded PDF file."""
    reader = PdfReader(pdf_file)
    text = ''
    for page in reader.pages:
        page_text = page.extract_text()
        if page_text:
            text += page_text + '\n'
    return text.strip()

def get_topic_from_text(text):
    """Use OpenAI to get the topic from the extracted text."""
    try:
        prompt_config = config['prompts']['topic_extraction']
        response = openai.ChatCompletion.create(
            model=OPENAI_MODEL,
            messages=[
                {
                    "role": "system",
                    "content": prompt_config['system_message']
                },
                {
                    "role": "user",
                    "content": f"{prompt_config['user_message']}\n\n{text}"
                }
            ],
            max_tokens=prompt_config['max_tokens'],
            temperature=prompt_config['temperature']
        )
        topic = response['choices'][0]['message']['content'].strip()
        return topic
    except Exception as e:
        print(f"Error getting topic from text: {e}")
        return None

def generate_webpage_content(topic):
    """Generate a detailed article based on the topic."""
    try:
        prompt_config = config['prompts']['content_generation']
        user_message = prompt_config['user_message'].replace("{topic}", topic)
        response = openai.ChatCompletion.create(
            model=OPENAI_MODEL,
            messages=[
                {
                    "role": "system",
                    "content": prompt_config['system_message']
                },
                {
                    "role": "user",
                    "content": user_message
                }
            ],
            max_tokens=prompt_config['max_tokens'],
            temperature=prompt_config['temperature']
        )
        content = response['choices'][0]['message']['content'].strip()
        return content
    except Exception as e:
        print(f"Error generating content: {e}")
        return None

def save_content_to_docx(content, filename):
    """Save the generated content into a DOCX file."""
    doc = Document()
    doc.add_heading('Generated Article', level=0)
    for line in content.split('\n'):
        if line.strip() == '':
            continue
        if line.startswith('#'):
            # Convert Markdown headings to Word headings
            heading_level = line.count('#')
            doc.add_heading(line.replace('#', '').strip(), level=heading_level)
        else:
            doc.add_paragraph(line)
    doc.save(filename)

@app.route('/', methods=['GET'])
def index():
    return render_template('index.html')

@app.route('/upload', methods=['POST'])
def upload_file():
    if 'file' not in request.files:
        return jsonify({'error': 'No file part in the request.'}), 400
    
    file = request.files['file']
    
    if file.filename == '':
        return jsonify({'error': 'No selected file.'}), 400
    
    if not file.filename.lower().endswith('.pdf'):
        return jsonify({'error': 'Invalid file type. Only PDF files are allowed.'}), 400
    
    # Extract text from PDF
    extracted_text = extract_text_from_pdf(file)
    if not extracted_text:
        return jsonify({'error': 'Failed to extract text from the PDF file.'}), 500
    
    # Get the topic from the text
    topic = get_topic_from_text(extracted_text)
    if not topic:
        return jsonify({'error': 'Failed to identify the topic from the extracted text.'}), 500
    
    # Generate the webpage content based on the topic
    webpage_content = generate_webpage_content(topic)
    if not webpage_content:
        return jsonify({'error': 'Failed to generate content based on the identified topic.'}), 500
    
    # Save the content to a DOCX file
    output_filename = "{}.docx".format(uuid.uuid4())
    save_content_to_docx(webpage_content, output_filename)
    
    return send_file(output_filename, as_attachment=True)

if __name__ == '__main__':
    app.run(
        host=config['app'].get('host', '127.0.0.1'),
        port=config['app'].get('port', 5000),
        debug=config['app'].get('debug', False)
    )
